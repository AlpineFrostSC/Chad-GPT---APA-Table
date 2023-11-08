from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement

# Create a new Word document
doc = Document()

# Add the table number and title
table_number = doc.add_paragraph()
run = table_number.add_run('Table 1')
run.bold = True

title = doc.add_paragraph()
run = title.add_run('Italicized Description of Table with Important Words Capitalized')
run.italic = True
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Simulated data based on the screenshot
headers = ["Independent Variables", "n", "min", "max", "mean", "Median", "SD", "Skew", "Kurtosis"]
rows = [
    ["NB mean scores", 185, 1.75, 7, 6.47, 6.75, 0.71, -2.906, 12.552],
    ["ESG mean scores", 184, 1.5, 7, 6.41, 6.63, 0.72, -2.836, 12.73],
    ["CBMCS mean scores", 185, 1.9, 3.95, 3.22, 3.19, 0.32, -0.178, 0.517],
    ["GICCS mean scores", 183, 3.9, 6.82, 5.63, 5.71, 0.77, -0.652, 0.085]
]

# Add the table to the Word document
table = doc.add_table(rows=1, cols=len(headers))

# Header row style
for i, header in enumerate(headers):
    cell = table.cell(0, i)
    cell.text = header
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.runs[0]
    run.font.bold = True

# Add the spanner row
spanner_row = table.add_row()
spanner_cell = spanner_row.cells[0]
# Merge all the cells in the spanner row
spanner_cell.merge(spanner_row.cells[-1])
spanner_cell.text = "Stage 1"
for paragraph in spanner_cell.paragraphs:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add the rest of the data to the table
for row_data in rows:
    row_cells = table.add_row().cells
    for i, cell_data in enumerate(row_data):
        row_cells[i].text = str(cell_data)
        paragraph = row_cells[i].paragraphs[0]
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Set the font for the entire table
for row in table.rows:
    for cell in row.cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.font.size = Pt(12)

# Remove all borders from the table
remove_table_borders(table)

# Save the document
doc_path = '/mnt/data/APA_Formatted_Document_With_Spanner_No_Borders.docx'
doc.save(doc_path)

doc_path
